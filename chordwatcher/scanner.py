"""Extract raw text out of the document types ChordWatcher can scan."""

from pathlib import Path


def extract_text(path) -> str:
    """Extract text from a .txt, .docx, or .pdf file.

    Raises ValueError for unsupported extensions and RuntimeError if a PDF
    turns out to be a scanned image that needs OCR dependencies we don't have.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return _extract_txt(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".doc":
        raise ValueError(
            "Legacy .doc files aren't supported. Please re-save the document as "
            ".docx or .txt and try again."
        )
    if suffix == ".pdf":
        return _extract_pdf(path)

    raise ValueError(
        f"Unsupported file type '{suffix}'. ChordWatcher can scan .txt, .docx, and .pdf files."
    )


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "Reading .docx files requires the 'python-docx' package. Install it with: "
            "pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "Reading .pdf files requires the 'pypdf' package. Install it with: pip install pypdf"
        ) from exc

    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)

    if len(text.strip()) < 20:
        # Little to no extractable text usually means this is a scanned image PDF.
        text = _ocr_pdf(path)

    return text


def _ocr_pdf(path: Path) -> str:
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError as exc:
        raise RuntimeError(
            f"'{path.name}' looks like a scanned PDF with no extractable text. Scanning it "
            "requires OCR support: pip install pdf2image pytesseract, plus the system tools "
            "'poppler' (for pdf2image) and 'tesseract-ocr' (for pytesseract)."
        ) from exc

    images = convert_from_path(str(path))
    return "\n".join(pytesseract.image_to_string(image) for image in images)
